import sys
import re
from docx import Document
from docx.shared import Inches
import os

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            continue

        # Handle Images
        img_match = re.search(r'!\[.*?\]\((.*?)\)', line)
        if img_match:
            img_path = img_match.group(1)
            # Resolve relative path
            full_img_path = os.path.join(os.path.dirname(md_path), img_path)
            if os.path.exists(full_img_path):
                doc.add_picture(full_img_path, width=Inches(5))
            else:
                doc.add_paragraph(f"[Image not found: {img_path}]")
            continue

        # Handle Tables (Simple | | version)
        if '|' in line:
            if not in_table:
                in_table = True
                table_data = []
            
            # Skip separator line | --- |
            if re.match(r'^[|\s-]*$', line):
                continue
                
            cells = [c.strip() for c in line.split('|') if c.strip() or line.split('|').index(c) not in [0, len(line.split('|'))-1]]
            if cells:
                table_data.append(cells)
            continue
        else:
            if in_table:
                # Flush table
                if table_data:
                    rows = len(table_data)
                    cols = max(len(r) for r in table_data)
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for i, r_data in enumerate(table_data):
                        for j, val in enumerate(r_data):
                            if j < cols:
                                table.cell(i, j).text = val
                in_table = False
                table_data = []

        # Handle text
        if line:
            # Handle Bold
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # Handle Links
            line = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', line)
            doc.add_paragraph(line)
        else:
            # Empty line for spacing
            pass

    # Save
    doc.save(docx_path)
    print(f"Saved to {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python convert_prd.py input.md output.docx")
    else:
        convert_md_to_docx(sys.argv[1], sys.argv[2])
